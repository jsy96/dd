# -*- coding: utf-8 -*-
"""
Vercel Serverless Function API
舱单数据处理应用
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import xlrd
import xlwt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import tempfile
from datetime import datetime
import base64

app = Flask(__name__)
CORS(app)


class ManifestProcessor:
    """舱单数据处理器"""

    def __init__(self, manifest_file_path):
        """初始化处理器"""
        self.manifest_data = self._read_manifest(manifest_file_path)

    def _read_manifest(self, file_path):
        """读取舱单Excel文件"""
        workbook = xlrd.open_workbook(file_path)
        sheet = workbook.sheet_by_index(0)

        data = {}
        for row_idx in range(sheet.nrows):
            row_values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]

            # 提取基本信息
            if '船名' in row_values[0]:
                data['vessel_name'] = row_values[1] if len(row_values) > 1 else ''
                data['voyage_no'] = row_values[4] if len(row_values) > 4 else ''
            elif '目的港' in row_values[0]:
                data['port_of_discharge'] = row_values[1] if len(row_values) > 1 else ''
            elif '总提单号' in row_values[0]:
                data['master_bl_no'] = row_values[1] if len(row_values) > 1 else ''

            # 提取分票统计数据
            elif '提单号' in row_values and '英文品名' in row_values:
                if row_idx + 1 < sheet.nrows:
                    data_row = [str(sheet.cell_value(row_idx + 1, col_idx)) for col_idx in range(sheet.ncols)]
                    data['bl_no'] = data_row[0]
                    data['cargo_name'] = data_row[2]
                    data['marks'] = data_row[5] if len(data_row) > 5 else 'N/M'
                    data['packages'] = data_row[9] if len(data_row) > 9 else ''
                    data['package_unit'] = data_row[10] if len(data_row) > 10 else 'CARTONS'
                    data['gross_weight'] = data_row[11] if len(data_row) > 11 else ''
                    data['volume'] = data_row[12] if len(data_row) > 12 else ''

            # 提取按箱统计数据
            elif '箱号' in row_values and '封号' in row_values:
                if row_idx + 1 < sheet.nrows:
                    data_row = [str(sheet.cell_value(row_idx + 1, col_idx)) for col_idx in range(sheet.ncols)]
                    data['container_no'] = data_row[0]
                    data['seal_no'] = data_row[1] if len(data_row) > 1 else ''
                    data['container_type'] = data_row[2] if len(data_row) > 2 else ''

        return data

    def generate_bl_confirmation(self, output_path, consignor_info=None, consignee_info=None, notify_party_info=None):
        """生成提单确认件"""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        def set_chinese_font(run, font_name='宋体', font_size=11):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

        p = doc.add_paragraph()
        run = p.add_run(f"船名航次：{self.manifest_data.get('vessel_name', '')} {self.manifest_data.get('voyage_no', '')}\t")
        set_chinese_font(run)
        run = p.add_run(f"目的港：{self.manifest_data.get('port_of_discharge', '')}\n")
        set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run(f"提单号：{self.manifest_data.get('bl_no', self.manifest_data.get('master_bl_no', ''))}\n")
        set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run(f"相应箱号 封号\n")
        set_chinese_font(run)
        run = p.add_run(f"箱号：{self.manifest_data.get('container_no', '')}\n")
        set_chinese_font(run)
        run = p.add_run(f"封号：{self.manifest_data.get('seal_no', '')}\n")
        set_chinese_font(run)
        run = p.add_run(f"箱型：{self.manifest_data.get('container_type', '')}\n")
        set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("发货人：\n")
        set_chinese_font(run)
        if consignor_info:
            for line in consignor_info.split('\n'):
                run = p.add_run(f"{line}\n")
                set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("收货人：\n")
        set_chinese_font(run)
        if consignee_info:
            for line in consignee_info.split('\n'):
                run = p.add_run(f"{line}\n")
                set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("通知人：\n")
        set_chinese_font(run)
        if notify_party_info:
            for line in notify_party_info.split('\n'):
                run = p.add_run(f"{line}\n")
                set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("品名：\n")
        set_chinese_font(run)
        cargo_names = self.manifest_data.get('cargo_name', '').split(',')
        for name in cargo_names:
            run = p.add_run(f"{name.strip()}\n")
            set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("件数／重量／体积\n")
        set_chinese_font(run)
        run = p.add_run(f"件数：{self.manifest_data.get('packages', '')} {self.manifest_data.get('package_unit', '')}\n")
        set_chinese_font(run)
        run = p.add_run(f"毛重：{self.manifest_data.get('gross_weight', '')} KGS\n")
        set_chinese_font(run)
        run = p.add_run(f"体积：{self.manifest_data.get('volume', '')} CBM\n")
        set_chinese_font(run)

        p = doc.add_paragraph()
        run = p.add_run("申请14天免用箱显示在提单上")
        set_chinese_font(run)

        doc.save(output_path)
        return output_path

    def generate_packing_list_invoice(self, output_path, invoice_no=None, invoice_date=None, consignee_name=None, items=None):
        """生成装箱单发票"""
        workbook = xlwt.Workbook(encoding='utf-8')
        sheet = workbook.add_sheet('装箱单发票')

        sheet.col(0).width = 256 * 15
        sheet.col(1).width = 256 * 8
        sheet.col(2).width = 256 * 10
        sheet.col(3).width = 256 * 35
        sheet.col(4).width = 256 * 10
        sheet.col(5).width = 256 * 8
        sheet.col(6).width = 256 * 10
        sheet.col(7).width = 256 * 10

        title_style = xlwt.XFStyle()
        title_font = xlwt.Font()
        title_font.height = 280
        title_style.font = title_font

        normal_style = xlwt.XFStyle()
        normal_font = xlwt.Font()
        normal_font.height = 220
        normal_style.font = normal_font

        bold_style = xlwt.XFStyle()
        bold_font = xlwt.Font()
        bold_font.height = 220
        bold_font.bold = True
        bold_style.font = bold_font

        sheet.write(1, 0, '浙江长江国际有限公司', title_style)
        sheet.write_merge(1, 1, 1, 14, '', normal_style)

        sheet.write(2, 0, 'ZHEJIANG CHEUNG KONG INTERNATIONAL LIMITED', normal_style)
        sheet.write_merge(2, 2, 1, 14, '', normal_style)

        sheet.write(3, 4, '发票', bold_style)
        sheet.write(3, 6, '第', normal_style)
        sheet.write(3, 7, invoice_no or 'YWSJ2602044', normal_style)
        sheet.write(3, 8, '号', normal_style)

        sheet.write(4, 5, 'No.', normal_style)
        sheet.write_merge(4, 4, 6, 8, '………………………………', normal_style)

        sheet.write(5, 4, 'INVOICE', bold_style)
        sheet.write(5, 6, '日期', normal_style)
        sheet.write(5, 7, invoice_date or datetime.now().strftime('%b.%d.%Y').upper(), normal_style)

        sheet.write(6, 6, 'Date……………………………………', normal_style)

        if consignee_name:
            sheet.write(7, 0, consignee_name, normal_style)
        else:
            sheet.write(7, 0, 'TO: SIJI SHIPPING L.L.C', normal_style)
        sheet.write_merge(7, 7, 1, 4, '', normal_style)
        sheet.write(7, 5, '信用证第', normal_style)
        sheet.write_merge(7, 7, 6, 8, '', normal_style)
        sheet.write(7, 9, '号', normal_style)

        sheet.write(8, 0, 'To:…………………………………………………………', normal_style)
        sheet.write(8, 5, 'L/C NO:………………………………', normal_style)

        sheet.write(9, 0, '唛头号码 Marks & Numbers', bold_style)
        sheet.write_merge(9, 9, 1, 2, '', normal_style)
        sheet.write(9, 3, '数量与品名 Quantities and Descriptions', bold_style)
        sheet.write_merge(9, 9, 4, 5, '', normal_style)
        sheet.write(9, 6, '单价 Unit price', bold_style)
        sheet.write_merge(9, 9, 7, 8, '金额 Amount', bold_style)

        marks = self.manifest_data.get('marks', 'N/M')
        sheet.write(10, 0, marks, normal_style)
        sheet.write_merge(10, 10, 1, 4, 'CIF DUBAI', normal_style)

        if items:
            row_idx = 11
            for item in items:
                sheet.write(row_idx, 1, str(item.get('qty', '')), normal_style)
                sheet.write(row_idx, 2, item.get('unit', 'CTNS'), normal_style)
                sheet.write(row_idx, 3, item.get('name', ''), normal_style)
                sheet.write(row_idx, 5, str(item.get('unit_price', '')), normal_style)
                sheet.write(row_idx, 6, '/CTNS', normal_style)
                sheet.write(row_idx, 7, str(item.get('amount', '')), normal_style)
                row_idx += 1

        workbook.save(output_path)
        return output_path


@app.route('/api/preview', methods=['POST', 'OPTIONS'])
def preview_data():
    """预览提取的数据"""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        if 'manifest_file' not in request.files:
            return jsonify({'error': '请上传舱单文件'}), 400

        file = request.files['manifest_file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        temp_manifest = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        file.save(temp_manifest.name)
        temp_manifest.close()

        processor = ManifestProcessor(temp_manifest.name)

        try:
            os.unlink(temp_manifest.name)
        except:
            pass

        return jsonify({
            'success': True,
            'data': processor.manifest_data
        })

    except Exception as e:
        return jsonify({'error': f'预览失败: {str(e)}'}), 500


@app.route('/api/process', methods=['POST', 'OPTIONS'])
def process_manifest():
    """处理舱单并生成文档"""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        if 'manifest_file' not in request.files:
            return jsonify({'error': '请上传舱单文件'}), 400

        file = request.files['manifest_file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        temp_manifest = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        file.save(temp_manifest.name)
        temp_manifest.close()

        invoice_no = request.form.get('invoice_no', '')
        invoice_date = request.form.get('invoice_date', '')
        consignor = request.form.get('consignor', '')
        consignee = request.form.get('consignee', '')
        notify_party = request.form.get('notify_party', '')

        items = []
        items_data = request.form.get('items', '')
        if items_data:
            for item_line in items_data.split('\n'):
                parts = item_line.split('|')
                if len(parts) >= 4:
                    items.append({
                        'qty': parts[0].strip(),
                        'unit': parts[1].strip(),
                        'name': parts[2].strip(),
                        'unit_price': parts[3].strip(),
                        'amount': parts[4].strip() if len(parts) > 4 else ''
                    })

        processor = ManifestProcessor(temp_manifest.name)

        bl_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        processor.generate_bl_confirmation(
            bl_output.name,
            consignor_info=consignor,
            consignee_info=consignee,
            notify_party_info=notify_party
        )

        pl_output = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        processor.generate_packing_list_invoice(
            pl_output.name,
            invoice_no=invoice_no,
            invoice_date=invoice_date,
            consignee_name=consignee.split('\n')[0] if consignee else None,
            items=items if items else None
        )

        # 读取文件并编码为 base64
        with open(bl_output.name, 'rb') as f:
            bl_data = base64.b64encode(f.read()).decode()
        with open(pl_output.name, 'rb') as f:
            pl_data = base64.b64encode(f.read()).decode()

        try:
            os.unlink(temp_manifest.name)
            os.unlink(bl_output.name)
            os.unlink(pl_output.name)
        except:
            pass

        return jsonify({
            'success': True,
            'message': '文档生成成功！',
            'bl_document': bl_data,
            'pl_document': pl_data,
            'extracted_data': processor.manifest_data
        })

    except Exception as e:
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


# Vercel 需要的入口点
def handler(event, context):
    """Vercel Serverless Function 入口点"""
    return app(event.get('body', ''), context)

# -*- coding: utf-8 -*-
"""
舱单数据处理应用
提取舱单数据并生成提单确认件和装箱单发票
"""

from flask import Flask, render_template, request, send_file, jsonify
from flask_cors import CORS
import xlrd
import xlwt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
import tempfile
import base64
from datetime import datetime

app = Flask(__name__)
CORS(app)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), kwargs[edge])
            tcBorders.append(edge_element)

    tcPr.append(tcBorders)
    return cell


class ManifestProcessor:
    """舱单数据处理器"""

    def __init__(self, manifest_file_path):
        """初始化处理器"""
        self.manifest_data = self._read_manifest(manifest_file_path)

    def _read_manifest(self, file_path):
        """读取舱单Excel文件"""
        workbook = xlrd.open_workbook(file_path)
        sheet = workbook.sheet_by_index(0)

        data = {}
        for row_idx in range(sheet.nrows):
            row_values = [str(sheet.cell_value(row_idx, col_idx)) if sheet.cell_value(row_idx, col_idx) else '' for col_idx in range(sheet.ncols)]

            # 1. 提取基本信息 (Row 3): 船名、航次、目的港
            if '船名' in row_values[0] and 'vessel_name' not in data:
                data['vessel_name'] = row_values[1] if len(row_values) > 1 else ''
                data['voyage_no'] = row_values[4] if len(row_values) > 4 else ''
                # 同时提取目的港（在同一行，查找"目的港"列）
                for col_idx, val in enumerate(row_values):
                    if val == '目的港' and col_idx + 1 < len(row_values) and 'port_of_discharge' not in data:
                        data['port_of_discharge'] = row_values[col_idx + 1]

            # 2. 提取总提单号 (Row 4)
            elif '总提单号' in row_values[0] and 'master_bl_no' not in data:
                data['master_bl_no'] = row_values[1] if len(row_values) > 1 else ''

            # 3. 提取分票统计数据 (Row 7-8) - 提单号、英文品名、唛头、件数、包装单位、毛重、体积
            elif ('提单号' in row_values and '英文品名' in row_values) and 'bl_no' not in data:
                if row_idx + 1 < sheet.nrows:
                    data_row = [str(sheet.cell_value(row_idx + 1, col_idx)) for col_idx in range(sheet.ncols)]
                    data['bl_no'] = data_row[0]
                    data['cargo_name'] = data_row[2]
                    data['marks'] = data_row[6] if len(data_row) > 6 else 'N/M'
                    data['packages'] = data_row[9] if len(data_row) > 9 else ''
                    data['package_unit'] = data_row[10] if len(data_row) > 10 else 'CARTONS'
                    data['gross_weight'] = data_row[11] if len(data_row) > 11 else ''
                    data['volume'] = data_row[12] if len(data_row) > 12 else ''

            # 4. 提取按箱统计数据 (Row 11-12) - 箱号、封号、箱型
            elif ('箱号' in row_values and '封号' in row_values and '提单号' in row_values) and 'container_no' not in data:
                if row_idx + 1 < sheet.nrows:
                    data_row = [str(sheet.cell_value(row_idx + 1, col_idx)) for col_idx in range(sheet.ncols)]
                    data['container_no'] = data_row[0]
                    data['seal_no'] = data_row[1] if len(data_row) > 1 else ''
                    data['container_type'] = data_row[2] if len(data_row) > 2 else ''

            # 5. 提取发货人信息 (Row 26-31)
            elif ('发货人' in ''.join(row_values) or 'Shipper' in ''.join(row_values)) and 'shipper' not in data:
                shipper = []
                for i in range(row_idx + 1, min(row_idx + 6, sheet.nrows)):
                    row_vals = [str(sheet.cell_value(i, col)) for col in range(sheet.ncols)]
                    # 第二列是标签（名称、地址等），第三列是值
                    if len(row_vals) > 2:
                        label = row_vals[1]
                        value = row_vals[2]
                        if '名称' in label and value:
                            shipper.append(value)
                        elif '地址' in label and value:
                            shipper.append(value)
                        elif '电话' in label and value:
                            shipper.append(f'TEL: {value}')
                data['shipper'] = '\\n'.join(shipper) if shipper else ''

            # 6. 提取收货人信息 (Row 33-40)
            elif ('收货人' in ''.join(row_values) or 'Consignee' in ''.join(row_values)) and 'consignee' not in data:
                consignee = []
                for i in range(row_idx + 1, min(row_idx + 8, sheet.nrows)):
                    row_vals = [str(sheet.cell_value(i, col)) for col in range(sheet.ncols)]
                    if len(row_vals) > 2:
                        label = row_vals[1]
                        value = row_vals[2]
                        if '名称' in label and value:
                            consignee.append(value)
                        elif '地址' in label and value:
                            consignee.append(value)
                        elif '电话' in label and value:
                            consignee.append(f'TEL: {value}')
                        elif '具体联系人' in label and value:
                            consignee.append(f'ATTN: {value}')
                        elif '联系人电话' in label and value:
                            consignee.append(f'MOB: {value}')
                data['consignee'] = '\\n'.join(consignee) if consignee else ''

            # 7. 提取通知人信息 (Row 42-47)
            elif ('通知人' in ''.join(row_values) or 'Notifier' in ''.join(row_values)) and 'notifier' not in data:
                notifier = []
                for i in range(row_idx + 1, min(row_idx + 6, sheet.nrows)):
                    row_vals = [str(sheet.cell_value(i, col)) for col in range(sheet.ncols)]
                    if len(row_vals) > 2:
                        label = row_vals[1]
                        value = row_vals[2]
                        if '名称' in label and value:
                            notifier.append(value)
                        elif '地址' in label and value:
                            notifier.append(value)
                        elif '电话' in label and value:
                            notifier.append(f'TEL: {value}')
                data['notifier'] = '\\n'.join(notifier) if notifier else ''

        return data

    def generate_bl_confirmation(self, output_path, consignor_info=None, consignee_info=None, notify_party_info=None):
        """生成提单确认件 - 完全使用舱单提取的数据"""
        doc = Document()

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        def add_run(paragraph, text, font_size=11, bold=False):
            run = paragraph.add_run(text)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(font_size)
            run.font.bold = bold
            return run

        # 船名航次和目的港
        p = doc.add_paragraph()
        add_run(p, f"船名航次：{self.manifest_data.get('vessel_name', '')} {self.manifest_data.get('voyage_no', '')}\t")
        add_run(p, f"目的港：{self.manifest_data.get('port_of_discharge', '')}\n")

        # 提单号
        p = doc.add_paragraph()
        add_run(p, f"提单号：{self.manifest_data.get('bl_no', self.manifest_data.get('master_bl_no', ''))}\n")

        # 箱号和封号
        p = doc.add_paragraph()
        add_run(p, f"相应箱号 封号\n")
        add_run(p, f"箱号：{self.manifest_data.get('container_no', '')}\n")
        add_run(p, f"封号：{self.manifest_data.get('seal_no', '')}\n")
        add_run(p, f"箱型：{self.manifest_data.get('container_type', '')}\n")

        # 发货人 - 完全使用舱单提取的数据
        p = doc.add_paragraph()
        add_run(p, "发货人：\n")
        shipper = self.manifest_data.get('shipper', '')
        if shipper:
            for line in shipper.split('\\n'):
                add_run(p, f"{line}\n")
        else:
            # 如果舱单中没有数据，使用用户输入（备用）
            if consignor_info:
                for line in consignor_info.split('\n'):
                    add_run(p, f"{line}\n")

        # 收货人 - 完全使用舱单提取的数据
        p = doc.add_paragraph()
        add_run(p, "收货人：\n")
        consignee = self.manifest_data.get('consignee', '')
        if consignee:
            for line in consignee.split('\\n'):
                add_run(p, f"{line}\n")
        else:
            # 如果舱单中没有数据，使用用户输入（备用）
            if consignee_info:
                for line in consignee_info.split('\n'):
                    add_run(p, f"{line}\n")

        # 通知人 - 完全使用舱单提取的数据
        p = doc.add_paragraph()
        add_run(p, "通知人：\n")
        notifier = self.manifest_data.get('notifier', '')
        if notifier:
            for line in notifier.split('\\n'):
                add_run(p, f"{line}\n")
        else:
            # 如果舱单中没有数据，使用用户输入（备用）
            if notify_party_info:
                for line in notify_party_info.split('\n'):
                    add_run(p, f"{line}\n")

        # 品名
        p = doc.add_paragraph()
        add_run(p, "品名：\n")
        cargo_names = self.manifest_data.get('cargo_name', '').split(',')
        for name in cargo_names:
            add_run(p, f"{name.strip()}\n")

        # 件数/重量/体积 - 使用舱单提取的数据
        p = doc.add_paragraph()
        add_run(p, "件数／重量／体积\n")
        add_run(p, f"件数：{self.manifest_data.get('packages', '')} {self.manifest_data.get('package_unit', '')}\n")
        add_run(p, f"毛重：{self.manifest_data.get('gross_weight', '')} KGS\n")
        add_run(p, f"体积：{self.manifest_data.get('volume', '')} CBM\n")

        # 免用箱申请
        p = doc.add_paragraph()
        add_run(p, "申请14天免用箱显示在提单上")

        doc.save(output_path)
        return output_path

    def generate_packing_list_invoice(self, output_path, invoice_no=None, invoice_date=None, consignee_name=None, items=None):
        """生成装箱单发票 - 完全按照格式要求"""
        workbook = xlwt.Workbook(encoding='utf-8')
        sheet = workbook.add_sheet('装箱单发票')

        # ========== 设置列宽 ==========
        # xlwt列宽单位：1单位 = 1/256字符宽
        # 原始数据：4192, 480, 1248, 1248, 8576, 2816, 1696, 1696, 1088, 1056
        sheet.col(0).width = 4192    # 约16.4字符
        sheet.col(1).width = 480     # 约1.9字符
        sheet.col(2).width = 1248    # 约4.9字符
        sheet.col(3).width = 1248    # 约4.9字符
        sheet.col(4).width = 8576    # 约33.5字符
        sheet.col(5).width = 2816    # 约11字符
        sheet.col(6).width = 1696    # 约6.6字符
        sheet.col(7).width = 1696    # 约6.6字符
        sheet.col(8).width = 1088    # 约4.3字符
        sheet.col(9).width = 1056    # 约4.1字符

        # ========== 定义样式函数 ==========
        def create_font(name, size, bold=False):
            f = xlwt.Font()
            f.name = name
            f.height = int(size * 20)  # 点数转twips
            f.bold = bold
            return f

        def create_borders(top=False, bottom=False, left=False, right=False):
            b = xlwt.Borders()
            if top:
                b.top = xlwt.Borders.THIN
            if bottom:
                b.bottom = xlwt.Borders.THIN
            if left:
                b.left = xlwt.Borders.THIN
            if right:
                b.right = xlwt.Borders.THIN
            return b

        def create_style(font, borders=None):
            s = xlwt.XFStyle()
            s.font = font
            if borders:
                s.borders = borders
            # 默认左对齐
            a = xlwt.Alignment()
            a.horz = xlwt.Alignment.HORZ_LEFT
            s.alignment = a
            return s

        # ========== 预定义样式 ==========
        # 中文标题样式（宋体24pt粗体）
        font_title_cn = create_font('宋体', 24, True)
        style_title_cn = create_style(font_title_cn)

        # 英文标题样式（宋体18pt粗体）
        font_title_en = create_font('宋体', 18, True)
        style_title_en = create_style(font_title_en)

        # 发票标题（宋体24pt）
        font_invoice = create_font('宋体', 24, False)
        style_invoice = create_style(font_invoice)

        # INVOICE（Times New Roman 16pt）
        font_invoice_en = create_font('Times New Roman', 16, False)
        style_invoice_en = create_style(font_invoice_en)

        # 普通文本（宋体10pt/12pt）
        font_normal_10 = create_font('宋体', 10, False)
        style_normal_10 = create_style(font_normal_10)

        font_normal_12 = create_font('宋体', 12, False)
        style_normal_12 = create_style(font_normal_12)

        # 数字文本（Times New Roman 11pt/12pt）
        font_number_11 = create_font('Times New Roman', 11, False)
        style_number_11 = create_style(font_number_11)

        font_number_12 = create_font('Times New Roman', 12, False)
        style_number_12 = create_style(font_number_12)

        # 表头样式（宋体10pt，带边框）
        font_header = create_font('宋体', 10, False)
        borders_all = create_borders(top=True, bottom=True, left=True, right=True)
        style_header = create_style(font_header, borders_all)

        # ========== 写入数据 ==========
        # Row 0: 空行
        pass

        # Row 1: 公司名称（宋体24pt粗体，合并C0-C9，行高465）
        sheet.row(1).height_mismatch = True
        sheet.row(1).height = 465
        sheet.write_merge(1, 1, 0, 9, '浙江长江国际有限公司', style_title_cn)

        # Row 2: 英文名称（宋体18pt粗体，合并C0-C9，行高420）
        sheet.row(2).height_mismatch = True
        sheet.row(2).height = 420
        sheet.write_merge(2, 2, 0, 9, '            ZHEJIANG CHEUNG KONG INTERNATIONAL LIMITED', style_title_en)

        # Row 3: 发票（宋体24pt，合并C4-C5，行高300）
        sheet.row(3).height_mismatch = True
        sheet.row(3).height = 300
        sheet.write_merge(3, 3, 4, 5, '发票', style_invoice)
        sheet.write(3, 6, '第', style_normal_10)
        # 发票号（合并C7-C9，宋体10pt）
        sheet.write_merge(3, 3, 7, 9, invoice_no or 'YWSJ2602044', style_normal_10)
        sheet.write(3, 10, '号', style_normal_10)

        # Row 4: No.（行高285）
        sheet.row(4).height_mismatch = True
        sheet.row(4).height = 285
        # 发票合并区域继续
        sheet.write_merge(4, 4, 4, 5, '', style_invoice)
        sheet.write(4, 6, 'No.', style_normal_10)
        sheet.write_merge(4, 4, 7, 9, '………………………………', style_normal_10)

        # Row 5: INVOICE（Times New Roman 16pt，合并C4-C5，行高300）
        sheet.row(5).height_mismatch = True
        sheet.row(5).height = 300
        sheet.write_merge(5, 5, 4, 5, 'INVOICE', style_invoice_en)
        sheet.write(5, 6, '日期', style_normal_10)
        sheet.write_merge(5, 5, 7, 9, invoice_date or datetime.now().strftime('%b.%d.%Y').upper(), style_normal_10)

        # Row 6: Date（行高345）
        sheet.row(6).height_mismatch = True
        sheet.row(6).height = 345
        sheet.write_merge(6, 6, 4, 5, '', style_invoice_en)
        sheet.write(6, 6, 'Date……………………………………', style_normal_12)

        # Row 7: 收货人（行高375）
        sheet.row(7).height_mismatch = True
        sheet.row(7).height = 375
        consignee = consignee_name or self.manifest_data.get('consignee', 'SIJI SHIPPING L.L.C')
        if '\\n' in consignee:
            consignee = consignee.split('\\n')[0]
        sheet.write(7, 0, consignee, style_normal_12)
        sheet.write(7, 6, '信用证第', style_normal_10)
        sheet.write(7, 10, '号', style_normal_10)

        # Row 8: To占位（行高345）
        sheet.row(8).height_mismatch = True
        sheet.row(8).height = 345
        sheet.write(8, 0, 'To:…………………………………………………………', style_normal_12)
        sheet.write(8, 6, 'L/C NO:……………………………', style_normal_10)

        # Row 9: 表头（行高510，带边框）
        sheet.row(9).height_mismatch = True
        sheet.row(9).height = 510
        sheet.write(9, 0, '唛头号码    Marks & Numbers', style_header)
        sheet.write_merge(9, 9, 2, 4, '数量与品名                                  Quantities and Descriptions', style_header)
        sheet.write_merge(9, 9, 5, 6, '单价            Unit price', style_header)
        sheet.write_merge(9, 9, 7, 9, '金额       Amount', style_header)

        # Row 10: N/M 和 CIF DUBAI（行高315，带边框）
        sheet.row(10).height_mismatch = True
        sheet.row(10).height = 315
        marks = self.manifest_data.get('marks', 'N/M')
        sheet.write(10, 0, marks, style_header)
        sheet.write_merge(10, 10, 5, 7, 'CIF DUBAI', style_header)

        # ========== 商品明细行 ==========
        if items:
            row_idx = 11
            for item in items:
                sheet.row(row_idx).height_mismatch = True
                sheet.row(row_idx).height = 315

                # C2: 数量
                sheet.write(row_idx, 2, str(item.get('qty', '')), style_number_12)

                # C3: 单位
                sheet.write(row_idx, 3, item.get('unit', 'CTNS'), style_normal_12)

                # C4: 品名
                sheet.write(row_idx, 4, item.get('name', ''), style_normal_12)

                # C5: 单价
                sheet.write(row_idx, 5, str(item.get('unit_price', '')), style_number_11)

                # C6: /CTNS
                sheet.write(row_idx, 6, '/CTNS', style_normal_12)

                # C7-C9: 金额（合并单元格）
                borders_sides = create_borders(left=True, right=True)
                style_amount = create_style(font_number_11, borders_sides)
                sheet.write_merge(row_idx, row_idx, 7, 9, str(item.get('amount', '')), style_amount)

                row_idx += 1

        workbook.save(output_path)
        return output_path


@app.route('/')
def index():
    """主页"""
    return render_template('index.html')


def file_to_base64(file_path):
    """将文件转换为 base64 编码"""
    with open(file_path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')


@app.route('/api/preview', methods=['POST', 'OPTIONS'])
def preview_data():
    """预览提取的数据"""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        if 'manifest_file' not in request.files:
            return jsonify({'error': '请上传舱单文件'}), 400

        file = request.files['manifest_file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        temp_manifest = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        file.save(temp_manifest.name)
        temp_manifest.close()

        processor = ManifestProcessor(temp_manifest.name)

        try:
            os.unlink(temp_manifest.name)
        except:
            pass

        return jsonify({
            'success': True,
            'data': processor.manifest_data
        })

    except Exception as e:
        return jsonify({'error': f'预览失败: {str(e)}'}), 500


@app.route('/api/process', methods=['POST', 'OPTIONS'])
def process_manifest():
    """处理舱单并生成文档"""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        if 'manifest_file' not in request.files:
            return jsonify({'error': '请上传舱单文件'}), 400

        file = request.files['manifest_file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        temp_manifest = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        file.save(temp_manifest.name)
        temp_manifest.close()

        invoice_no = request.form.get('invoice_no', '')
        invoice_date = request.form.get('invoice_date', '')
        consignor = request.form.get('consignor', '')
        consignee = request.form.get('consignee', '')
        notify_party = request.form.get('notify_party', '')

        items = []
        items_data = request.form.get('items', '')
        if items_data:
            for item_line in items_data.split('\n'):
                parts = item_line.split('|')
                if len(parts) >= 4:
                    items.append({
                        'qty': parts[0].strip(),
                        'unit': parts[1].strip(),
                        'name': parts[2].strip(),
                        'unit_price': parts[3].strip(),
                        'amount': parts[4].strip() if len(parts) > 4 else ''
                    })

        processor = ManifestProcessor(temp_manifest.name)

        bl_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        processor.generate_bl_confirmation(
            bl_output.name,
            consignor_info=consignor,
            consignee_info=consignee,
            notify_party_info=notify_party
        )

        pl_output = tempfile.NamedTemporaryFile(delete=False, suffix='.xls')
        processor.generate_packing_list_invoice(
            pl_output.name,
            invoice_no=invoice_no,
            invoice_date=invoice_date,
            consignee_name=consignee.split('\n')[0] if consignee else None,
            items=items if items else None
        )

        bl_data = file_to_base64(bl_output.name)
        pl_data = file_to_base64(pl_output.name)

        try:
            os.unlink(temp_manifest.name)
            os.unlink(bl_output.name)
            os.unlink(pl_output.name)
        except:
            pass

        return jsonify({
            'success': True,
            'message': '文档生成成功！',
            'bl_document': bl_data,
            'pl_document': pl_data,
            'extracted_data': processor.manifest_data
        })

    except Exception as e:
        import traceback
        return jsonify({'error': f'处理失败: {str(e)}\n{traceback.format_exc()}'}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
